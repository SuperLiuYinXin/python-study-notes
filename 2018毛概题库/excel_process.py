# coding=utf-8
import xlrd
import os
import re
# pip install python-docx ! 不是docx
from docx import Document



# 我的毛概文件存放目录
xl_dir = './毛概试题库excel'

# 指定生成的docx的文件夹
# 每个人的不一样，按照需求指定
# 我的电脑的Ubuntu,和windows的路径不一样，按需制定
doc_dir = '/home/liuyinxin/Documents/2018毛概试题库/毛概试题库DOC'

# 获取排序所需要的key
def arr_key(x):
    m = re.match('(\d+)\..*', x)
    if m is not None:
        return int(m.group(1))
    return 1


# 　获取题库文件excel存放的目录下的所有文件
xl_list = [os.path.join(xl_dir, _) for _ in sorted(os.listdir(xl_dir), key=arr_key)]

# 确保输出目录存在
if not os.path.exists(doc_dir):
    os.makedirs(doc_dir)

# 题目分类
S = 1 # 单选题
M = 2 # 多选题
F = 3 # 填空题
J = 4 # 判断题
SA = 5 #简答题



#　题目类型转换,通过　'单选题'，等获取题目的数字编号
# 方便处理
q_type_to_int = {
    '单选题': S,
    '多选题': M,
    '填空题': F,
    '判断题': J,
    '简答题': SA
}

# 从题目数字编号变成文字
q_type_to_str = ['', '单选', '多选', '填空', '判断', '简答']


# 定义题目类
class Question:

    # 构造函数
    def __init__(self, q_type, content, ans, choose_len, chooses):
        """
        :param q_type:  题目分类
        :param content: 问题
        :param ans: 答案
        :param choose_len: 选项长度
        :param chooses: 选项
        """
        self.content = content.strip()
        self.type = q_type_to_int[q_type]
        # 如果是简答题或者是填空题
        if self.type == SA or self.type == F:
            self.ans = ' '.join(chooses)
        # 如果是判断题，答案只有　A 和　B ,获取ascii编码相减,获取正确答案
        elif self.type == J:
            self.ans = chooses[ord(ans) - ord('A')]
        else:
            self.ans = ans
            self.choose_len = choose_len
            self.chooses = chooses

    def __str__(self):
        return self.content + '题目类型:' + q_type_to_str[self.type]




# 生成文档
def gen_doc(doc, header, quests, with_ans=True, choose_end='\t', is_break=False, filter=[S, M, J, M, SA]):
    """
    :param doc: 文档对象，如果是空，则创建
    :param header: 文档标题
    :param quests: 所有的问题
    :param with_ans: 是否显示答案
    :param choose_end: 每个选项结尾的符号，默认为制表符
    :param is_break: 在文档最后是否加一个分页
    :param filter: 过滤要显示的题,只生成过滤器里的
    :return: 返回文档对象
    """
    # 如果是空,就新建一个文档
    if doc is None:
        doc = Document()

    # 向文档中添加标题
    doc.add_heading(header)
    for i, quest in enumerate(quests):
        # 如果题目的类型不在要求的类型之中，就不生成它
        if quest.type not in filter:
            continue
        # 格式化标题
        title = "{}. [{}] {} " \
            .format(i + 1, q_type_to_str[quest.type], quest.content)
        # 添加段落, 并且不为简单题
        if with_ans and quest.type != SA:
            title += ' ' + "({})".format(quest.ans)

        # 将标题添加到文档中
        title_p = doc.add_paragraph(title)

        #  处理选项
        choose_str = ""
        #  按照题目类型分类，如果是判断题或者简答,或者填空题，没有选项
        if quest.type <= M:
            for i, choose in enumerate(quest.chooses):
                choose_str += "{}. {}{}".format(chr(ord('A') + i), choose, choose_end)
            choose_p = doc.add_paragraph(choose_str)
        # 如果是简答题
        elif quest.type == SA:
            choose_str = '答: ' + quest.ans
            choose_p = doc.add_paragraph(choose_str)

    # 是否加入分页符号
    if is_break:
        doc.add_page_break()
    return doc

# 主函数
if __name__ == '__main__':
    all_result = {}
    # 根据路径获取路径内的所有文件读取
    # 这里假设输入文件下所有的文件都是要读取的excel题库
    for xl_path in xl_list:
        # 通过xlrd打开excel库
        with xlrd.open_workbook(xl_path) as book:
            questions = []
            # 获取第一个excel表
            sheet = book.sheet_by_index(0)
            # 获取excel表的所有行
            lines = sheet.nrows
            for l in range(1, lines):
                # 获取题目的类型，　读取的是字符串　如'简答题'，　'单选题' 等
                q_type = sheet.cell(l, 1).value
                content = sheet.cell(l, 2).value #　获取题目内容
                ans = sheet.cell(l, 3).value # 获取答案
                choose_len = int(sheet.cell(l, 6).value) # 获取答案的长度
                # 依次获取答案里的内容,主要要用下str都变成字符串
                chooses = [str(sheet.cell(l, _).value) for _ in range(7, 7 + choose_len)]
                # 将获取的信息
                q = Question(q_type, content, ans, choose_len, chooses)
                # 添加到题目列表里
                questions.append(q)
            # 按照题的类型排序
            questions = sorted(questions, key=lambda i: i.type)
            # 这时候的questions是一个excel文件里的所有题，按照文件名写入全部题的字典
            # 以文件名作为　key , 以　题作为value
            all_result[os.path.basename(xl_path)] = questions

    # 新建一个word文档
    doc = Document()
    # 将所有的写入到doc文件中
    for k, quests in all_result.items():
        title = k[:k.rfind('.')]  # 标题
        # 生成文件
        path = os.path.join(doc_dir, title)
        print('finish[', path, ']')

        # 将题目和标题写入doc文档中
        gen_doc(doc, title, quests, with_ans=True, choose_end='\t', filter=[S, M])
    # 保存doc文档,按照类型自己起个名字
    doc.save(os.path.join(doc_dir, '2018_毛概题库_8章后_单选多选_有答案_打印版(页数更少).docx'))
